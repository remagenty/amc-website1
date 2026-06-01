import json
import os
from docx import Document

with open(os.path.expanduser('~/Downloads/catalogue_promove.json'), 'r', encoding='utf-8') as f:
    data = json.load(f)

doc = Document()
doc.add_heading('URLS DES OUTILS PROMOVE DEMOLITION', 0)
doc.add_paragraph('Cliquez sur chaque lien pour telecharger les images.')

by_cat = {}
for m in data['machines']:
    c = m.get('categorie', 'Autre')
    if c not in by_cat:
        by_cat[c] = []
    by_cat[c].append(m)

for cat in sorted(by_cat.keys()):
    doc.add_heading(cat, level=1)
    for m in by_cat[cat]:
        doc.add_paragraph(m['marque'] + ' ' + m['modele'], style='Heading 3')
        url = m.get('url_promove_demolition') or m.get('url_promove')
        if url:
            doc.add_paragraph(url)
        doc.add_paragraph()

doc.save(os.path.expanduser('~/Downloads/URLs_Promove.docx'))
print('OK: URLs_Promove.docx')
