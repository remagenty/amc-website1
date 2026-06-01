import json
import os
from docx import Document

with open(os.path.expanduser('~/Downloads/catalogue_magni.json'), 'r', encoding='utf-8') as f:
    data = json.load(f)

doc = Document()
doc.add_heading('URLS DES MACHINES MAGNI', 0)
doc.add_paragraph('Cliquez sur chaque lien pour telecharger les images.')

by_cat = {}
for m in data['machines']:
    c = m.get('categorie', 'Autre')
    if c not in by_cat:
        by_cat[c] = []
    by_cat[c].append(m)

for cat in sorted(by_cat.keys()):
    doc.add_heading(cat, level=1)
    for m in by_cat[cat]:
        doc.add_paragraph(m['marque'] + ' ' + m['modele'], style='Heading 3')
        if m.get('url_magni'):
            doc.add_paragraph(m['url_magni'])
        doc.add_paragraph()

doc.save(os.path.expanduser('~/Downloads/URLs_Magni.docx'))
print('OK: URLs_Magni.docx')
